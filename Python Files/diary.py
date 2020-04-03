import datetime
import time
import os
import ez
from docx import Document
from docx.shared import Pt, RGBColor, Inches
leapYear = lambda year: ((year % 4 == 0 and year % 100 != 0) or year % 400 == 0)
os.chdir('C:\\Users\\Seaky\\Desktop\\Docs\\日记')

def main(startYear = 0, endYear = 0):
    d = ['Mon', '一', 'Tue', '二', 'Wed', '三', 'Thu', '四', 'Fri', '五', 'Sat', '六', 'Sun', '天', '年0', '年', '月0', '月']
    startYear = startYear or datetime.date.today().year
    endYear = endYear or startYear
    for year in range(startYear, endYear + 1):
        days = 365 + leapYear(year)
        t = datetime.datetime(year, 1, 1).timestamp()
        s = '\n'.join([time.strftime('%Y年%m月%d日   星期%a   天气：\n\n', time.localtime(t + 86400 * day)) for day in range(days)])
        s = ez.sub(s, *d)
        document = Document()
        paragraph = document.add_paragraph(s)
        document.styles['Normal'].paragraph_format.line_spacing = 1.15
        document.styles['Normal'].paragraph_format.space_after = Pt(0)
        document.styles['Normal'].font.size = Pt(10.5)
        document.styles['Normal'].font.name = u'宋体'
        document.save(f'{year}.docx')
        print(f'{year} Done.')
        
def earlyDiary(startYear = 0, endYear = 0):
    d = ['Mon', '一', 'Tue', '二', 'Wed', '三', 'Thu', '四', 'Fri', '五', 'Sat', '六', 'Sun', '天', '年0', '年', '月0', '月']
    startYear = startYear or datetime.date.today().year
    endYear = endYear or startYear
    for year in range(startYear, endYear + 1):
        days = 365 + leapYear(year)
        t = datetime.datetime(year, 1, 1).timestamp()
        document = Document()
        for day in range(132):
            s = ez.sub(time.strftime('%Y年%m月%d日   星期%a   天气：   今天的心情：', time.localtime(t + 86400 * day)), *d)
            paragraph = document.add_paragraph(s)
            paragraph.style.paragraph_format.line_spacing = 1.15
            paragraph.style.paragraph_format.space_after = Pt(0)
            paragraph.style.font.size = Pt(10.5)
            paragraph.style.font.name = u'宋体'
            table = document.add_table(9, 2, 'Table Grid')
            row0 = table.rows[0]
            row0.cells[0].merge(row0.cells[1])
            for row in table.rows[-4:-1]:
                row.cells[0].merge(row.cells[1])
            for row, text in zip(table.rows, \
                                 ['', '今日好事', '今天的反省', '明天的目标', '今日特讯', '今天领到的钱：元', '今天花掉的钱：元', '今天剩余的钱：元', '明天的预算']):
                row.cells[0].text = text
            table.cell(2, 1).text = 'Waste time'
            table.cell(3, 1).text = 'Much better'
            for row in table.rows[1:]:
                row.cells[0].width = Inches(1)
##            for row in table.rows:
##                for cell in row.cells:
##                    style = cell.paragraphs[0].style
##                    style.font.size = Pt(10.5)
##                    style.font.name = u'宋体'
##                    cell.paragraphs[0].style = style
            document.add_paragraph('')
##        paragraph.style.font.color = RGBColor(0xFF, 0xFF, 0xFF)

        document.save(f'早期日记：{year}.docx')
        print(f'{year} Done.')

def earlyDiary2(startYear = 0, endYear = 0):
    d = ['Mon', '一', 'Tue', '二', 'Wed', '三', 'Thu', '四', 'Fri', '五', 'Sat', '六', 'Sun', '天', '年0', '年', '月0', '月']
    startYear = startYear or datetime.date.today().year
    endYear = endYear or startYear
    for year in range(startYear, endYear + 1):
        days = 365 + leapYear(year)
        t = datetime.datetime(year, 1, 1).timestamp()
        document = Document()
        for day in range(2):
            s = ez.sub(time.strftime('%Y年%m月%d日   星期%a   天气：   今天的心情：', time.localtime(t + 86400 * day)), *d)
            paragraph = document.add_paragraph(s)
            paragraph.style.paragraph_format.line_spacing = 1.15
            paragraph.style.paragraph_format.space_after = Pt(0)
            paragraph.style.font.size = Pt(10.5)
            paragraph.style.font.name = u'宋体'
            table = document.add_table(9, 2, 'Table Grid')
            table.cell(0, 0).merge(table.cell(0, 1))
            table.cell(2, 0).merge(table.cell(3, 0)).merge(table.cell(4, 0))
            table.cell(7, 1).merge(table.cell(8, 1))
            for row in table.rows[1:]:
                row.cells[0].width = Inches(2)
            for (x, y), text in zip([(1, 0), (1, 1), (3, 1), (5, 0), (6, 0), (6, 1), (7, 0), (8, 0), (9, 0)],\
                                    ['今日好事', '今天的反省', '明天的目标', '今日特讯', '今天领到的钱：元', '明天的预算', '今天花掉的钱：元', '今天剩余的钱：元']):
                table.cell(x, y).text = text
            document.add_paragraph('')
##        paragraph.style.font.color = RGBColor(0xFF, 0xFF, 0xFF)

        document.save(f'早期日记：{year}.docx')
        print(f'{year} Done.')

def search(word, groupByYear = False, Print = True, printText = True):
    results = {}
    date = None
    for doc in os.listdir():
        try:
            d = Document(doc)
        except:
            continue
        for paragraph in d.paragraphs:
            for text in paragraph.text.split('\n'):
                if all(flag in text for flag in ['年', '月', '日', '星期']):
                    date = text.split()[0]
                    if groupByYear:
                        date = date.split('年')[0]
                elif word in text:
                    lst = results.get(date, [])
                    lst.append(text)
                    results[date] = lst
    if Print:
        count = 0
        for date, lst in results.items():
            print(f'{date}：{len(lst)}')
            count += len(lst)
            if printText:
                for text in lst:
                    print(text)
                print()
        print(f'共{count}条结果')
    return results

if __name__ == '__main__':
##    main(2021)
    search('崇仁', 0, 1, 1)
